from pathlib import Path
from uuid import uuid4

from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_ollama import OllamaEmbeddings
from langchain_qdrant import QdrantVectorStore
from pypdf import PdfReader
from qdrant_client import QdrantClient
from qdrant_client.http.models import Distance, VectorParams

from app.config import get_settings


settings = get_settings()


def get_embeddings() -> OllamaEmbeddings:
    return OllamaEmbeddings(
        model=settings.ollama_embed_model,
        base_url=settings.ollama_base_url,
    )


def ensure_collection() -> None:
    client = QdrantClient(url=settings.qdrant_url)
    collections = client.get_collections().collections
    if any(item.name == settings.qdrant_collection for item in collections):
        return

    vector_size = len(get_embeddings().embed_query("dimension probe"))
    client.create_collection(
        collection_name=settings.qdrant_collection,
        vectors_config=VectorParams(size=vector_size, distance=Distance.COSINE),
    )


def get_vector_store() -> QdrantVectorStore:
    ensure_collection()
    return QdrantVectorStore(
        client=QdrantClient(url=settings.qdrant_url),
        collection_name=settings.qdrant_collection,
        embedding=get_embeddings(),
    )


def read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[Page {index}]\n{text}")
    return "\n\n".join(pages)


def read_docx(path: Path) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())


def read_image_ocr(path: Path) -> str:
    try:
        import pytesseract
        from PIL import Image
    except ImportError as exc:
        raise RuntimeError("Install pillow, pytesseract, and Tesseract OCR to ingest images.") from exc

    return pytesseract.image_to_string(Image.open(path))


def load_file_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return read_pdf(path)
    if suffix == ".docx":
        return read_docx(path)
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix in {".png", ".jpg", ".jpeg"}:
        return read_image_ocr(path)
    raise ValueError(f"Unsupported file type: {path.suffix}")


def chunk_documents(title: str, text: str, source_type: str, source_path: str | None = None) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(chunk_size=900, chunk_overlap=150)
    docs = [
        Document(
            page_content=chunk,
            metadata={
                "title": title,
                "source_type": source_type,
                "source_path": source_path or title,
            },
        )
        for chunk in splitter.split_text(text)
        if chunk.strip()
    ]
    return docs


def ingest_text(title: str, text: str, source_type: str = "manual") -> int:
    docs = chunk_documents(title=title, text=text, source_type=source_type)
    if not docs:
        return 0
    store = get_vector_store()
    store.add_documents(docs, ids=[str(uuid4()) for _ in docs])
    return len(docs)


def ingest_file(path: Path, source_type: str = "file") -> int:
    text = load_file_text(path)
    docs = chunk_documents(
        title=path.name,
        text=text,
        source_type=source_type,
        source_path=str(path.resolve()),
    )
    if not docs:
        return 0
    store = get_vector_store()
    store.add_documents(docs, ids=[str(uuid4()) for _ in docs])
    return len(docs)


def retrieve_documents(question: str, k: int = 5) -> list[dict[str, str]]:
    store = get_vector_store()
    docs = store.similarity_search(question, k=k)
    return [
        {
            "content": doc.page_content,
            "title": str(doc.metadata.get("title", "")),
            "source_type": str(doc.metadata.get("source_type", "")),
            "source_path": str(doc.metadata.get("source_path", "")),
        }
        for doc in docs
    ]

